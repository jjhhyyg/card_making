from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt
from docx2pdf import convert
import random

sign = 'Erikkson'


def createPdf(wordPath, pdfPath):
    """
    word转pdf
    :param wordPath: word文件路径
    :param pdfPath: 生成pdf文件路径
    """
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordPath, ReadOnly=1)
    doc.ExportAsFixedFormat(pdfPath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)


def read_txt(file_name):
    """
    读入txt文件
    :param file_name: 文件名（带后缀）
    :return: 以行返回列表
    """
    result = []
    with open(file_name, encoding='utf-8') as file:
        for line in file.readlines():
            line = line.rstrip()
            result.append(line)
    return result


if __name__ == '__main__':
    name_list = read_txt('names.txt')
    sentence_list = read_txt('sentences.txt')
    card_name_list = []

    for i in range(len(name_list)):
        document = Document()
        # 获取本文档中的所有章节
        sections = document.sections
        # 将该章节中的纸张方向设置为横向
        for section in sections:
            # 需要同时设置width,height才能成功
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height

        # 文档字体设置
        document.styles['Normal'].font.name = u'华文行楷'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文行楷')

        name = name_list[i]
        random_index = random.randint(0, len(sentence_list) - 1)
        sentence = sentence_list[random_index]
        card_name = f"{sign}给{name}的祝福"
        card_name_list.append(card_name)

        # 名称左对齐
        para = document.add_paragraph()
        run = para.add_run(name + '：')
        run.font.size = Pt(16)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.space_after = Pt(18)

        # 祝福语首行缩进2字符
        para = document.add_paragraph()
        run = para.add_run(sentence)
        run.font.size = Pt(20)
        run.font.name = 'Bradley Hand ITC'
        para.paragraph_format.first_line_indent = 406400

        # 署名右对齐
        para = document.add_paragraph()
        run = para.add_run('--' + sign)
        run.font.name = 'Segoe Script'
        run.font.size = Pt(16)
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        para.space_before = Pt(18)

        document.save(card_name + '.docx')

    for card_name in card_name_list:
        # 转换为pdf文档
        name_docx = card_name + '.docx'
        name_pdf = card_name + '.pdf'
        convert(name_docx, name_pdf)
